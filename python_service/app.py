# E:/solux-maker/python_service/app.py

from flask import Flask, request, jsonify
import os
import uuid
import docling as dl
import json
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'

# Garantir que as pastas existam
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "ok", "service": "docling-processor"})

# Adicionando uma rota raiz para evitar 404
@app.route('/', methods=['GET'])
def index():
    return jsonify({
        "service": "docling-processor",
        "status": "running",
        "endpoints": ["/", "/health", "/process"]
    })

@app.route('/process', methods=['POST'])
def process_document():
    if 'file' not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado"}), 400
    
    file = request.files['file']
    document_id = request.form.get('document_id', str(uuid.uuid4()))
    document_type = request.form.get('document_type', '').lower()
    
    if file.filename == '':
        return jsonify({"error": "Nome de arquivo vazio"}), 400
    
    # Salvar o arquivo
    file_path = os.path.join(UPLOAD_FOLDER, f"{document_id}.{document_type}")
    file.save(file_path)
    
    try:
        # Print para debug
        print(f"Processando arquivo: {file_path}, tipo: {document_type}")
        
        # Extrair texto usando métodos diferentes baseados no tipo de arquivo
        extracted_text = ""
        metadata = {"title": file.filename}
        
        # Processamento baseado no tipo de documento
        if document_type.lower() == 'pdf':
            try:
                # Método 1: Tentar com dl.parse
                doc = dl.parse(file_path)
                extracted_text = doc.text()
                # Adicionar metadados disponíveis
                metadata.update({
                    "method": "docling.parse",
                    "type": "pdf"
                })
            except Exception as e1:
                print(f"Erro com dl.parse para PDF: {e1}")
                # Método 2: Usar fallback para PDFs
                import PyPDF2
                try:
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        extracted_text = ""
                        for page in reader.pages:
                            extracted_text += page.extract_text() + "\n\n"
                        metadata.update({
                            "method": "PyPDF2",
                            "type": "pdf",
                            "pages": len(reader.pages)
                        })
                except Exception as e2:
                    print(f"Fallback para PDF também falhou: {e2}")
                    raise e2
            
        elif document_type.lower() == 'docx':
            try:
                # Método 1: Tentar com dl.parse
                doc = dl.parse(file_path)
                extracted_text = doc.text()
                metadata.update({
                    "method": "docling.parse",
                    "type": "docx"
                })
            except Exception as e1:
                print(f"Erro com dl.parse para DOCX: {e1}")
                # Método 2: Usar docx diretamente
                import docx
                try:
                    doc = docx.Document(file_path)
                    extracted_text = "\n\n".join([para.text for para in doc.paragraphs])
                    metadata.update({
                        "method": "python-docx",
                        "type": "docx",
                        "paragraphs": len(doc.paragraphs)
                    })
                except Exception as e2:
                    print(f"Fallback para DOCX também falhou: {e2}")
                    raise e2
            
        elif document_type.lower() == 'txt':
            # Para TXT, lemos diretamente
            with open(file_path, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
            metadata.update({
                "method": "direct_read",
                "type": "txt"
            })
            
        else:
            return jsonify({
                "error": f"Tipo de documento não suportado: {document_type}",
                "document_id": document_id
            }), 400
        
        # Verificar se extraímos texto
        if not extracted_text:
            return jsonify({
                "error": "Não foi possível extrair texto do documento",
                "document_id": document_id
            }), 500
        
        # Adicionar metadados adicionais
        metadata.update({
            "word_count": len(extracted_text.split()),
            "char_count": len(extracted_text)
        })
        
        # Salvar o resultado processado
        result_path = os.path.join(PROCESSED_FOLDER, f"{document_id}.json")
        result_data = {
            "document_id": document_id,
            "text": extracted_text,
            "metadata": metadata
        }
        
        with open(result_path, 'w', encoding='utf-8') as f:
            json.dump(result_data, f, ensure_ascii=False)
        
        # Print para debug
        print(f"Processamento concluído. Extraídos {metadata['char_count']} caracteres.")
        
        # Retornar resultado completo
        return jsonify({
            "success": True,
            "document_id": document_id,
            "metadata": metadata,
            "text": extracted_text
        })
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Erro ao processar documento: {error_details}")
        
        # Retornar erro com detalhes
        return jsonify({
            "success": False,  # Explicitamente marcado como falha
            "error": str(e),
            "traceback": error_details,
            "document_id": document_id
        }), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)